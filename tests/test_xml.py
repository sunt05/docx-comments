"""Word Online compatibility and XML structure tests."""

from docx import Document

from docx_comments import CommentManager, PersonInfo


def author_obj(name: str) -> PersonInfo:
    return PersonInfo(author=name)


class TestWordOnlineCompatibility:
    """Tests for Word Online compatibility by validating XML structure."""

    def test_xml_parts_created(self, tmp_path):
        """Test that all required XML parts are created."""
        from zipfile import ZipFile

        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)

        # Add a comment
        mgr.add_comment(para, "Test comment", author_obj("Author"))

        # Save document
        output_path = tmp_path / "test_parts.docx"
        doc.save(str(output_path))

        # Check ZIP contents
        with ZipFile(str(output_path), "r") as zf:
            names = zf.namelist()
            assert "word/comments.xml" in names
            assert "word/commentsExtended.xml" in names
            assert "word/commentsIds.xml" in names

    def test_comments_xml_structure(self, tmp_path):
        """Test comments.xml has correct structure."""
        from lxml import etree
        from zipfile import ZipFile

        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)

        comment_id = mgr.add_comment(
            para, "Review this section", author_obj("Reviewer"), initials="R"
        )

        output_path = tmp_path / "test_comments_xml.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            xml = etree.fromstring(zf.read("word/comments.xml"))

        # Check root element
        assert xml.tag.endswith("}comments")

        # Find comment element
        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        comments = xml.findall(f"{{{ns_w}}}comment")
        assert len(comments) == 1

        comment = comments[0]
        assert comment.get(f"{{{ns_w}}}author") == "Reviewer"
        assert comment.get(f"{{{ns_w}}}initials") == "R"
        assert comment.get(f"{{{ns_w}}}id") == comment_id

    def test_threading_xml_structure(self, tmp_path):
        """Test commentsExtended.xml has correct threading structure."""
        from lxml import etree
        from zipfile import ZipFile

        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)

        # Add root comment and reply
        root_id = mgr.add_comment(para, "Root comment", author_obj("Author1"))
        mgr.reply_to_comment(root_id, "Reply comment", author_obj("Author2"))

        output_path = tmp_path / "test_threading.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            xml = etree.fromstring(zf.read("word/commentsExtended.xml"))

        # Check root element
        assert xml.tag.endswith("}commentsEx")

        # Find commentEx elements
        ns_w15 = "http://schemas.microsoft.com/office/word/2012/wordml"
        comment_exs = xml.findall(f"{{{ns_w15}}}commentEx")
        assert len(comment_exs) == 2

        # Check that reply has parent link
        para_ids = set()
        parent_links = {}
        for ce in comment_exs:
            para_id = ce.get(f"{{{ns_w15}}}paraId")
            parent = ce.get(f"{{{ns_w15}}}paraIdParent")
            para_ids.add(para_id)
            if parent:
                parent_links[para_id] = parent

        # One comment should have a parent link
        assert len(parent_links) == 1
        # The parent should exist
        assert list(parent_links.values())[0] in para_ids

    def test_resolved_status_in_xml(self, tmp_path):
        """Test that resolved status is correctly saved in XML."""
        from lxml import etree
        from zipfile import ZipFile

        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)

        comment_id = mgr.add_comment(para, "Comment to resolve", author_obj("Author"))
        mgr.resolve_comment(comment_id)

        output_path = tmp_path / "test_resolved.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            xml = etree.fromstring(zf.read("word/commentsExtended.xml"))

        ns_w15 = "http://schemas.microsoft.com/office/word/2012/wordml"
        comment_exs = xml.findall(f"{{{ns_w15}}}commentEx")
        assert len(comment_exs) == 1
        assert comment_exs[0].get(f"{{{ns_w15}}}done") == "1"

    def test_document_xml_anchors(self, tmp_path):
        """Test that document.xml has proper comment anchors."""
        from lxml import etree
        from zipfile import ZipFile

        doc = Document()
        para = doc.add_paragraph("This is test text to comment on.")
        mgr = CommentManager(doc)

        comment_id = mgr.add_comment(para, "Test comment", author_obj("Author"))

        output_path = tmp_path / "test_anchors.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            xml = etree.fromstring(zf.read("word/document.xml"))

        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        # Find comment anchors
        range_start = xml.find(f".//{{{ns_w}}}commentRangeStart")
        range_end = xml.find(f".//{{{ns_w}}}commentRangeEnd")
        comment_ref = xml.find(f".//{{{ns_w}}}commentReference")

        assert range_start is not None, "commentRangeStart not found"
        assert range_end is not None, "commentRangeEnd not found"
        assert comment_ref is not None, "commentReference not found"

        # Verify IDs match
        assert range_start.get(f"{{{ns_w}}}id") == comment_id
        assert range_end.get(f"{{{ns_w}}}id") == comment_id
        assert comment_ref.get(f"{{{ns_w}}}id") == comment_id

    def test_reply_anchor_ordering(self, tmp_path):
        """Ensure reply anchors keep commentRangeEnd before commentReference runs."""
        from lxml import etree
        from zipfile import ZipFile

        doc = Document()
        para = doc.add_paragraph("This is test text to comment on.")
        mgr = CommentManager(doc)

        root_id = mgr.add_comment(para, "Root comment", author_obj("Author"))
        mgr.reply_to_comment(root_id, "Reply comment", author_obj("Author2"))

        output_path = tmp_path / "test_reply_anchor_order.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            xml = etree.fromstring(zf.read("word/document.xml"))

        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        range_start = xml.find(
            f".//{{{ns_w}}}commentRangeStart[@{{{ns_w}}}id='{root_id}']"
        )
        assert range_start is not None

        para_elem = range_start.getparent()
        children = list(para_elem)

        end_indices = [
            idx
            for idx, child in enumerate(children)
            if etree.QName(child).localname == "commentRangeEnd"
        ]
        ref_indices = [
            idx
            for idx, child in enumerate(children)
            if etree.QName(child).localname == "r"
            and child.find(f"{{{ns_w}}}commentReference") is not None
        ]

        assert end_indices, "commentRangeEnd elements not found"
        assert ref_indices, "commentReference runs not found"
        assert max(end_indices) < min(ref_indices)

    def test_full_roundtrip(self, tmp_path):
        """Test full save/reload roundtrip with all features."""
        doc = Document()
        para1 = doc.add_paragraph("First paragraph for comments")
        para2 = doc.add_paragraph("Second paragraph for comments")
        mgr = CommentManager(doc)

        # Add various comments
        id1 = mgr.add_comment(para1, "Comment on first para", author_obj("Alice"), "A")
        id2 = mgr.add_comment(para2, "Comment on second para", author_obj("Bob"), "B")
        mgr.reply_to_comment(id1, "Reply to Alice", author_obj("Charlie"), "C")
        mgr.resolve_comment(id2)

        # Save
        output_path = tmp_path / "test_roundtrip.docx"
        doc.save(str(output_path))

        # Reload
        doc2 = Document(str(output_path))
        mgr2 = CommentManager(doc2)

        # Verify comments
        comments = list(mgr2.list_comments())
        assert len(comments) == 3

        # Verify threads
        threads = mgr2.get_comment_threads()
        assert len(threads) == 2

        # Find thread with reply
        thread_with_reply = next(t for t in threads if t.reply_count > 0)
        assert thread_with_reply.root.text == "Comment on first para"
        assert thread_with_reply.replies[0].text == "Reply to Alice"

        # Find resolved thread
        resolved_thread = next(t for t in threads if t.root.text == "Comment on second para")
        assert resolved_thread.is_resolved
