"""Tests for metadata migration."""

from docx import Document

from docx_comments import CommentManager, PersonInfo


def author_obj(name: str) -> PersonInfo:
    return PersonInfo(author=name)


class TestCommentMigration:
    """Tests for comment metadata migration."""

    def test_migrate_comment_metadata(self):
        """Test backfilling missing comment metadata."""
        from docx_comments.xml_parts import CommentsExtendedPart, CommentsIdsPart

        doc = Document()
        para = doc.add_paragraph("Text with comment")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "Comment", author_obj("Author"))

        # Remove paraId/textId from comments.xml
        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        ns_w14 = "http://schemas.microsoft.com/office/word/2010/wordml"
        comment_elem = mgr._comments_xml.find(f"{{{ns_w}}}comment")
        para_elem = comment_elem.find(f"{{{ns_w}}}p")
        para_elem.attrib.pop(f"{{{ns_w14}}}paraId", None)
        para_elem.attrib.pop(f"{{{ns_w14}}}textId", None)
        mgr._save_comments()

        # Clear threading and durable IDs
        ext_part = CommentsExtendedPart(doc)
        ids_part = CommentsIdsPart(doc)
        for elem in list(ext_part.xml):
            ext_part.xml.remove(elem)
        ext_part._save()
        for elem in list(ids_part.xml):
            ids_part.xml.remove(elem)
        ids_part._save()

        # Run migration
        mgr.migrate_comment_metadata()

        # Verify metadata restored
        para_id = para_elem.get(f"{{{ns_w14}}}paraId")
        text_id = para_elem.get(f"{{{ns_w14}}}textId")
        assert para_id is not None
        assert text_id is not None
        assert para_id in CommentsExtendedPart(doc).get_threading_info()
        assert para_id in CommentsIdsPart(doc).get_durable_ids()

    def test_migrate_comment_extensible_date_utc(self):
        """Test backfilling missing dateUtc in commentsExtensible.xml."""
        from lxml import etree

        from docx_comments.xml_parts import CommentsExtensiblePart, CommentsIdsPart

        doc = Document()
        para = doc.add_paragraph("Text with comment")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "Comment", author_obj("Author"))

        ids_part = CommentsIdsPart(doc)
        durable_ids = ids_part.get_durable_ids()
        assert durable_ids
        durable_id = next(iter(durable_ids.values()))

        ext_part = CommentsExtensiblePart(doc)
        ns_w16cex = "http://schemas.microsoft.com/office/word/2018/wordml/cex"
        date_attr = f"{{{ns_w16cex}}}dateUtc"
        durable_attr = f"{{{ns_w16cex}}}durableId"

        removed = False
        for elem in ext_part.xml:
            if (
                etree.QName(elem).localname == "commentExtensible"
                and elem.get(durable_attr) == durable_id
            ):
                if date_attr in elem.attrib:
                    elem.attrib.pop(date_attr, None)
                    removed = True
                break
        assert removed
        ext_part._save()

        # Run migration and ensure dateUtc is restored.
        mgr.migrate_comment_metadata()

        ext_part = CommentsExtensiblePart(doc)
        for elem in ext_part.xml:
            if (
                etree.QName(elem).localname == "commentExtensible"
                and elem.get(durable_attr) == durable_id
            ):
                assert elem.get(date_attr) is not None
                break
        else:
            raise AssertionError("commentExtensible entry not found")
