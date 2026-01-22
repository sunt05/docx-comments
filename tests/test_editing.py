"""Tests for editing comment anchors and lifecycle operations."""

from docx import Document

from docx_comments import CommentManager, PersonInfo
from docx_comments.anchors import CommentAnchor
from docx_comments.xml_parts import CommentsExtendedPart, CommentsIdsPart


def author_obj(name: str) -> PersonInfo:
    return PersonInfo(author=name)


class TestCommentEditing:
    """Tests for comment deletion and re-anchoring behavior."""

    def test_unresolve_comment(self):
        """Resolved comments can be marked unresolved."""
        doc = Document()
        para = doc.add_paragraph("Test paragraph")
        mgr = CommentManager(doc)

        comment_id = mgr.add_comment(
            paragraph=para,
            text="Needs work",
            author=author_obj("Reviewer"),
        )

        mgr.resolve_comment(comment_id)
        mgr.unresolve_comment(comment_id)

        comments = list(mgr.list_comments())
        assert len(comments) == 1
        assert not comments[0].is_resolved

    def test_delete_comment_detaches_replies(self):
        """Deleting a root comment detaches remaining replies."""
        doc = Document()
        para = doc.add_paragraph("Test paragraph")
        mgr = CommentManager(doc)

        root_id = mgr.add_comment(
            paragraph=para,
            text="Root comment",
            author=author_obj("Author1"),
        )
        reply_id = mgr.reply_to_comment(
            parent_id=root_id,
            text="Reply comment",
            author=author_obj("Author2"),
        )

        mgr.delete_comment(root_id)

        comments = list(mgr.list_comments())
        assert len(comments) == 1
        assert comments[0].comment_id == reply_id
        assert comments[0].parent_para_id is None

    def test_delete_thread_removes_all(self):
        """Deleting a thread removes root and replies."""
        doc = Document()
        para = doc.add_paragraph("Test paragraph")
        mgr = CommentManager(doc)

        root_id = mgr.add_comment(
            paragraph=para,
            text="Root comment",
            author=author_obj("Author1"),
        )
        reply_id = mgr.reply_to_comment(
            parent_id=root_id,
            text="Reply comment",
            author=author_obj("Author2"),
        )

        mgr.delete_thread(reply_id)

        comments = list(mgr.list_comments())
        assert len(comments) == 0

        anchor = CommentAnchor(doc)
        assert anchor.find_paragraph_with_comment(root_id) is None
        assert anchor.find_paragraph_with_comment(reply_id) is None

    def test_delete_comment_cleans_orphan_metadata(self):
        """Deleting a comment cleans orphan metadata and detaches replies."""
        doc = Document()
        para = doc.add_paragraph("Test paragraph")
        mgr = CommentManager(doc)

        root_id = mgr.add_comment(
            paragraph=para,
            text="Root comment",
            author=author_obj("Author1"),
        )
        reply_id = mgr.reply_to_comment(
            parent_id=root_id,
            text="Reply comment",
            author=author_obj("Author2"),
        )

        root_para_id = next(
            comment.para_id
            for comment in mgr.list_comments()
            if comment.comment_id == root_id
        )

        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        ns_w14 = "http://schemas.microsoft.com/office/word/2010/wordml"
        for comment_elem in mgr._comments_xml.findall(f"{{{ns_w}}}comment"):
            if comment_elem.get(f"{{{ns_w}}}id") == root_id:
                para_elem = comment_elem.find(f"{{{ns_w}}}p")
                para_elem.attrib.pop(f"{{{ns_w14}}}paraId", None)
                para_elem.attrib.pop(f"{{{ns_w14}}}textId", None)
                break
        mgr._save_comments()

        assert root_para_id in CommentsExtendedPart(doc).get_threading_info()
        assert root_para_id in CommentsIdsPart(doc).get_durable_ids()

        mgr.delete_comment(root_id)

        comments = list(mgr.list_comments())
        assert len(comments) == 1
        assert comments[0].comment_id == reply_id
        assert comments[0].parent_para_id is None

        assert root_para_id not in CommentsExtendedPart(doc).get_threading_info()
        assert root_para_id not in CommentsIdsPart(doc).get_durable_ids()

    def test_move_comment_updates_anchor_paragraph(self):
        """Moving a comment updates its anchor location."""
        doc = Document()
        para1 = doc.add_paragraph("Paragraph one")
        para2 = doc.add_paragraph("Paragraph two")
        mgr = CommentManager(doc)

        comment_id = mgr.add_comment(
            paragraph=para1,
            text="Move me",
            author=author_obj("Author1"),
        )

        mgr.move_comment(comment_id, para2)

        anchor = CommentAnchor(doc)
        anchored_para = anchor.find_paragraph_with_comment(comment_id)
        assert anchored_para is not None
        assert anchored_para._element is para2._element

    def test_move_thread_moves_replies(self):
        """Moving a thread re-anchors replies at the new location."""
        doc = Document()
        para1 = doc.add_paragraph("Paragraph one")
        para2 = doc.add_paragraph("Paragraph two")
        mgr = CommentManager(doc)

        root_id = mgr.add_comment(
            paragraph=para1,
            text="Root comment",
            author=author_obj("Author1"),
        )
        reply_id = mgr.reply_to_comment(
            parent_id=root_id,
            text="Reply comment",
            author=author_obj("Author2"),
        )

        mgr.move_thread(root_id, para2)

        anchor = CommentAnchor(doc)
        anchored_root = anchor.find_paragraph_with_comment(root_id)
        anchored_reply = anchor.find_paragraph_with_comment(reply_id)
        assert anchored_root is not None
        assert anchored_reply is not None
        assert anchored_root._element is para2._element
        assert anchored_reply._element is para2._element
