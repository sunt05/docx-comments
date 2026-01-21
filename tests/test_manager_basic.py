"""Basic CommentManager behavior tests."""

from docx import Document

import pytest

from docx_comments import CommentManager, PersonInfo


def author_obj(name: str) -> PersonInfo:
    return PersonInfo(author=name)


class TestCommentManagerBasic:
    """Tests for core CommentManager behavior."""

    def test_init_new_document(self):
        """Test initializing with a new document."""
        doc = Document()
        doc.add_paragraph("Test paragraph")
        mgr = CommentManager(doc)
        assert mgr is not None

    def test_list_comments_empty(self):
        """Test listing comments on document with no comments."""
        doc = Document()
        doc.add_paragraph("Test paragraph")
        mgr = CommentManager(doc)
        comments = list(mgr.list_comments())
        assert len(comments) == 0

    def test_add_comment(self, tmp_path):
        """Test adding a comment to a paragraph."""
        doc = Document()
        para = doc.add_paragraph("This is test text to comment on.")
        mgr = CommentManager(doc)

        comment_id = mgr.add_comment(
            paragraph=para,
            text="This is a test comment",
            author=author_obj("Test Author"),
            initials="TA",
        )

        assert comment_id is not None
        assert len(comment_id) > 0

        # Verify comment was added
        comments = list(mgr.list_comments())
        assert len(comments) == 1
        assert comments[0].text == "This is a test comment"
        assert comments[0].author == "Test Author"
        assert comments[0].initials == "TA"

        # Save and reload
        output_path = tmp_path / "test_output.docx"
        doc.save(str(output_path))

        # Reload and verify
        doc2 = Document(str(output_path))
        mgr2 = CommentManager(doc2)
        comments2 = list(mgr2.list_comments())
        assert len(comments2) == 1

    def test_add_comment_rejects_non_personinfo_author(self):
        """Author must be PersonInfo."""
        doc = Document()
        para = doc.add_paragraph("This is test text to comment on.")
        mgr = CommentManager(doc)

        with pytest.raises(TypeError):
            mgr.add_comment(
                paragraph=para,
                text="This is a test comment",
                author="Test Author",
            )

        with pytest.raises(TypeError):
            mgr.add_comment(
                paragraph=para,
                text="This is a test comment",
                author={"author": "Test Author"},
            )

    def test_resolve_comment(self):
        """Test marking a comment as resolved."""
        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)

        comment_id = mgr.add_comment(
            paragraph=para,
            text="Comment to resolve",
            author=author_obj("Author"),
        )

        # Initially not resolved
        comments = list(mgr.list_comments())
        assert not comments[0].is_resolved

        # Resolve
        mgr.resolve_comment(comment_id)

        # Verify resolved
        comments = list(mgr.list_comments())
        assert comments[0].is_resolved

    def test_get_comment_threads(self):
        """Test getting comment threads."""
        doc = Document()
        para1 = doc.add_paragraph("First paragraph")
        para2 = doc.add_paragraph("Second paragraph")
        mgr = CommentManager(doc)

        # Add two independent comments
        id1 = mgr.add_comment(para1, "Comment 1", author_obj("Author1"))
        id2 = mgr.add_comment(para2, "Comment 2", author_obj("Author2"))

        # Add replies to first comment
        mgr.reply_to_comment(id1, "Reply 1a", author_obj("Author3"))
        mgr.reply_to_comment(id1, "Reply 1b", author_obj("Author4"))

        threads = mgr.get_comment_threads()
        assert len(threads) == 2

        # Find thread with replies
        thread_with_replies = next(t for t in threads if t.reply_count > 0)
        assert thread_with_replies.root.text == "Comment 1"
        assert thread_with_replies.reply_count == 2
