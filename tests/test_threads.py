"""Threading and reply behavior tests."""

from docx import Document

from docx_comments import CommentManager, PersonInfo


def author_obj(name: str) -> PersonInfo:
    return PersonInfo(author=name)


class TestCommentThreads:
    """Tests for threaded comment behavior."""

    def test_reply_to_comment(self, tmp_path):
        """Test replying to a comment."""
        doc = Document()
        para = doc.add_paragraph("Text to comment on.")
        mgr = CommentManager(doc)

        # Add root comment
        root_id = mgr.add_comment(
            paragraph=para,
            text="Root comment",
            author=author_obj("Author1"),
        )

        # Add reply
        reply_id = mgr.reply_to_comment(
            parent_id=root_id,
            text="Reply comment",
            author=author_obj("Author2"),
        )

        assert reply_id != root_id

        # Check threading
        threads = mgr.get_comment_threads()
        assert len(threads) == 1
        assert threads[0].root.text == "Root comment"
        assert len(threads[0].replies) == 1
        assert threads[0].replies[0].text == "Reply comment"

    def test_reply_to_reply(self):
        """Test replying to a reply in a thread."""
        doc = Document()
        para = doc.add_paragraph("Threaded comment text.")
        mgr = CommentManager(doc)

        root_id = mgr.add_comment(para, "Root comment", author_obj("Author1"))
        reply1_id = mgr.reply_to_comment(root_id, "Reply 1", author_obj("Author2"))
        reply2_id = mgr.reply_to_comment(reply1_id, "Reply 2", author_obj("Author3"))

        comments = list(mgr.list_comments())
        reply2 = next(c for c in comments if c.comment_id == reply2_id)
        root = next(c for c in comments if c.comment_id == root_id)
        assert reply2.parent_para_id == root.para_id

        threads = mgr.get_comment_threads()
        assert len(threads) == 1
        assert threads[0].root.comment_id == root_id
        assert threads[0].reply_count == 2

    def test_reply_to_comment_in_table(self):
        """Test replying to a comment anchored in a table."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell_para = table.rows[0].cells[0].paragraphs[0]
        cell_para.add_run("Cell text")
        mgr = CommentManager(doc)

        root_id = mgr.add_comment(cell_para, "Table comment", author_obj("Author1"))
        reply_id = mgr.reply_to_comment(root_id, "Reply to table", author_obj("Author2"))

        assert reply_id != root_id

    def test_reply_to_comment_in_header(self):
        """Test replying to a comment anchored in a header."""
        doc = Document()
        section = doc.sections[0]
        header_para = section.header.paragraphs[0]
        header_para.add_run("Header text")
        mgr = CommentManager(doc)

        root_id = mgr.add_comment(header_para, "Header comment", author_obj("Author1"))
        reply_id = mgr.reply_to_comment(root_id, "Reply to header", author_obj("Author2"))

        assert reply_id != root_id

    def test_multiple_threads_multiple_replies_roundtrip(self, tmp_path):
        """Test replies across multiple threads and locations with roundtrip."""
        doc = Document()
        para1 = doc.add_paragraph("Body para 1")
        para2 = doc.add_paragraph("Body para 2")
        table = doc.add_table(rows=1, cols=1)
        table_para = table.rows[0].cells[0].paragraphs[0]
        table_para.add_run("Table text")
        header_para = doc.sections[0].header.paragraphs[0]
        header_para.add_run("Header text")

        mgr = CommentManager(doc)

        root_body = mgr.add_comment(para1, "Root body", author_obj("Author1"))
        root_table = mgr.add_comment(table_para, "Root table", author_obj("Author2"))
        root_header = mgr.add_comment(header_para, "Root header", author_obj("Author3"))

        reply_body_1 = mgr.reply_to_comment(root_body, "Body reply 1", author_obj("Author4"))
        reply_body_2 = mgr.reply_to_comment(reply_body_1, "Body reply 2", author_obj("Author5"))
        mgr.reply_to_comment(root_table, "Table reply 1", author_obj("Author6"))
        mgr.reply_to_comment(root_header, "Header reply 1", author_obj("Author7"))

        output_path = tmp_path / "multi_thread_roundtrip.docx"
        doc.save(str(output_path))

        doc2 = Document(str(output_path))
        mgr2 = CommentManager(doc2)

        threads = mgr2.get_comment_threads()
        assert len(threads) == 3

        by_root = {t.root.text: t for t in threads}
        assert by_root["Root body"].reply_count == 2
        assert by_root["Root table"].reply_count == 1
        assert by_root["Root header"].reply_count == 1

        comments = list(mgr2.list_comments())
        reply2 = next(c for c in comments if c.text == "Body reply 2")
        root = next(c for c in comments if c.text == "Root body")
        assert reply2.parent_para_id == root.para_id
