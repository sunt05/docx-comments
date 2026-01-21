"""Basic tests for docx-comments module."""

import pytest
from docx import Document

from docx_comments import CommentManager, PersonInfo


def author_obj(name: str) -> PersonInfo:
    return PersonInfo(author=name)


class TestCommentManager:
    """Tests for CommentManager class."""

    def test_init_new_document(self):
        """Test initializing with a new document."""
        doc = Document()
        doc.add_paragraph("Test paragraph")
        mgr = CommentManager(doc)
        assert mgr is not None

    def test_list_comments_empty(self):
        """Test listing comments on document with no comments."""
        doc = Document()
        doc.add_paragraph("Test paragraph")
        mgr = CommentManager(doc)
        comments = list(mgr.list_comments())
        assert len(comments) == 0

    def test_add_comment(self, tmp_path):
        """Test adding a comment to a paragraph."""
        doc = Document()
        para = doc.add_paragraph("This is test text to comment on.")
        mgr = CommentManager(doc)

        comment_id = mgr.add_comment(
            paragraph=para,
            text="This is a test comment",
            author=author_obj("Test Author"),
            initials="TA",
        )

        assert comment_id is not None
        assert len(comment_id) > 0

        # Verify comment was added
        comments = list(mgr.list_comments())
        assert len(comments) == 1
        assert comments[0].text == "This is a test comment"
        assert comments[0].author == "Test Author"
        assert comments[0].initials == "TA"

        # Save and reload
        output_path = tmp_path / "test_output.docx"
        doc.save(str(output_path))

        # Reload and verify
        doc2 = Document(str(output_path))
        mgr2 = CommentManager(doc2)
        comments2 = list(mgr2.list_comments())
        assert len(comments2) == 1

    def test_add_comment_rejects_string_author(self):
        """Author string should be rejected to avoid ambiguous identity."""
        doc = Document()
        para = doc.add_paragraph("This is test text to comment on.")
        mgr = CommentManager(doc)

        with pytest.raises(TypeError):
            mgr.add_comment(
                paragraph=para,
                text="This is a test comment",
                author="Test Author",
            )

        with pytest.raises(TypeError):
            mgr.add_comment(
                paragraph=para,
                text="This is a test comment",
                author={"author": "Test Author"},
            )

    def test_reply_to_comment(self, tmp_path):
        """Test replying to a comment."""
        doc = Document()
        para = doc.add_paragraph("Text to comment on.")
        mgr = CommentManager(doc)

        # Add root comment
        root_id = mgr.add_comment(
            paragraph=para,
            text="Root comment",
            author=author_obj("Author1"),
        )

        # Add reply
        reply_id = mgr.reply_to_comment(
            parent_id=root_id,
            text="Reply comment",
            author=author_obj("Author2"),
        )

        assert reply_id != root_id

        # Check threading
        threads = mgr.get_comment_threads()
        assert len(threads) == 1
        assert threads[0].root.text == "Root comment"
        assert len(threads[0].replies) == 1
        assert threads[0].replies[0].text == "Reply comment"

    def test_resolve_comment(self):
        """Test marking a comment as resolved."""
        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)

        comment_id = mgr.add_comment(
            paragraph=para,
            text="Comment to resolve",
            author=author_obj("Author"),
        )

        # Initially not resolved
        comments = list(mgr.list_comments())
        assert not comments[0].is_resolved

        # Resolve
        mgr.resolve_comment(comment_id)

        # Verify resolved
        comments = list(mgr.list_comments())
        assert comments[0].is_resolved

    def test_get_comment_threads(self):
        """Test getting comment threads."""
        doc = Document()
        para1 = doc.add_paragraph("First paragraph")
        para2 = doc.add_paragraph("Second paragraph")
        mgr = CommentManager(doc)

        # Add two independent comments
        id1 = mgr.add_comment(para1, "Comment 1", author_obj("Author1"))
        id2 = mgr.add_comment(para2, "Comment 2", author_obj("Author2"))

        # Add replies to first comment
        mgr.reply_to_comment(id1, "Reply 1a", author_obj("Author3"))
        mgr.reply_to_comment(id1, "Reply 1b", author_obj("Author4"))

        threads = mgr.get_comment_threads()
        assert len(threads) == 2

        # Find thread with replies
        thread_with_replies = next(t for t in threads if t.reply_count > 0)
        assert thread_with_replies.root.text == "Comment 1"
        assert thread_with_replies.reply_count == 2


class TestCommentInfo:
    """Tests for CommentInfo model."""

    def test_is_reply(self):
        """Test is_reply property."""
        from docx_comments.models import CommentInfo

        root = CommentInfo(
            comment_id="1",
            para_id="ABC",
            text="Root",
            author="Author",
        )
        assert not root.is_reply

        reply = CommentInfo(
            comment_id="2",
            para_id="DEF",
            text="Reply",
            author="Author",
            parent_para_id="ABC",
        )
        assert reply.is_reply


class TestCommentThread:
    """Tests for CommentThread model."""

    def test_all_comments(self):
        """Test all_comments property."""
        from docx_comments.models import CommentInfo, CommentThread

        root = CommentInfo("1", "ABC", "Root", "Author")
        reply1 = CommentInfo("2", "DEF", "Reply1", "Author", parent_para_id="ABC")
        reply2 = CommentInfo("3", "GHI", "Reply2", "Author", parent_para_id="ABC")

        thread = CommentThread(root=root, replies=[reply1, reply2])

        assert len(thread.all_comments) == 3
        assert thread.all_comments[0] is root


class TestWordOnlineCompatibility:
    """Tests for Word Online compatibility by validating XML structure."""

    def test_xml_parts_created(self, tmp_path):
        """Test that all required XML parts are created."""
        from zipfile import ZipFile

        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)

        # Add a comment
        mgr.add_comment(para, "Test comment", author_obj("Author"))

        # Save document
        output_path = tmp_path / "test_parts.docx"
        doc.save(str(output_path))

        # Check ZIP contents
        with ZipFile(str(output_path), "r") as zf:
            names = zf.namelist()
            assert "word/comments.xml" in names
            assert "word/commentsExtended.xml" in names
            assert "word/commentsIds.xml" in names

    def test_comments_xml_structure(self, tmp_path):
        """Test comments.xml has correct structure."""
        from lxml import etree
        from zipfile import ZipFile

        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)

        comment_id = mgr.add_comment(
            para, "Review this section", author_obj("Reviewer"), initials="R"
        )

        output_path = tmp_path / "test_comments_xml.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            xml = etree.fromstring(zf.read("word/comments.xml"))

        # Check root element
        assert xml.tag.endswith("}comments")

        # Find comment element
        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        comments = xml.findall(f"{{{ns_w}}}comment")
        assert len(comments) == 1

        comment = comments[0]
        assert comment.get(f"{{{ns_w}}}author") == "Reviewer"
        assert comment.get(f"{{{ns_w}}}initials") == "R"
        assert comment.get(f"{{{ns_w}}}id") == comment_id

    def test_threading_xml_structure(self, tmp_path):
        """Test commentsExtended.xml has correct threading structure."""
        from lxml import etree
        from zipfile import ZipFile

        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)

        # Add root comment and reply
        root_id = mgr.add_comment(para, "Root comment", author_obj("Author1"))
        reply_id = mgr.reply_to_comment(root_id, "Reply comment", author_obj("Author2"))

        output_path = tmp_path / "test_threading.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            xml = etree.fromstring(zf.read("word/commentsExtended.xml"))

        # Check root element
        assert xml.tag.endswith("}commentsEx")

        # Find commentEx elements
        ns_w15 = "http://schemas.microsoft.com/office/word/2012/wordml"
        comment_exs = xml.findall(f"{{{ns_w15}}}commentEx")
        assert len(comment_exs) == 2

        # Check that reply has parent link
        para_ids = set()
        parent_links = {}
        for ce in comment_exs:
            para_id = ce.get(f"{{{ns_w15}}}paraId")
            parent = ce.get(f"{{{ns_w15}}}paraIdParent")
            para_ids.add(para_id)
            if parent:
                parent_links[para_id] = parent

        # One comment should have a parent link
        assert len(parent_links) == 1
        # The parent should exist
        assert list(parent_links.values())[0] in para_ids

    def test_resolved_status_in_xml(self, tmp_path):
        """Test that resolved status is correctly saved in XML."""
        from lxml import etree
        from zipfile import ZipFile

        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)

        comment_id = mgr.add_comment(para, "Comment to resolve", author_obj("Author"))
        mgr.resolve_comment(comment_id)

        output_path = tmp_path / "test_resolved.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            xml = etree.fromstring(zf.read("word/commentsExtended.xml"))

        ns_w15 = "http://schemas.microsoft.com/office/word/2012/wordml"
        comment_exs = xml.findall(f"{{{ns_w15}}}commentEx")
        assert len(comment_exs) == 1
        assert comment_exs[0].get(f"{{{ns_w15}}}done") == "1"

    def test_document_xml_anchors(self, tmp_path):
        """Test that document.xml has proper comment anchors."""
        from lxml import etree
        from zipfile import ZipFile

        doc = Document()
        para = doc.add_paragraph("This is test text to comment on.")
        mgr = CommentManager(doc)

        comment_id = mgr.add_comment(para, "Test comment", author_obj("Author"))

        output_path = tmp_path / "test_anchors.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            xml = etree.fromstring(zf.read("word/document.xml"))

        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        # Find comment anchors
        range_start = xml.find(f".//{{{ns_w}}}commentRangeStart")
        range_end = xml.find(f".//{{{ns_w}}}commentRangeEnd")
        comment_ref = xml.find(f".//{{{ns_w}}}commentReference")

        assert range_start is not None, "commentRangeStart not found"
        assert range_end is not None, "commentRangeEnd not found"
        assert comment_ref is not None, "commentReference not found"

        # Verify IDs match
        assert range_start.get(f"{{{ns_w}}}id") == comment_id
        assert range_end.get(f"{{{ns_w}}}id") == comment_id
        assert comment_ref.get(f"{{{ns_w}}}id") == comment_id

    def test_full_roundtrip(self, tmp_path):
        """Test full save/reload roundtrip with all features."""
        doc = Document()
        para1 = doc.add_paragraph("First paragraph for comments")
        para2 = doc.add_paragraph("Second paragraph for comments")
        mgr = CommentManager(doc)

        # Add various comments
        id1 = mgr.add_comment(para1, "Comment on first para", author_obj("Alice"), "A")
        id2 = mgr.add_comment(para2, "Comment on second para", author_obj("Bob"), "B")
        id3 = mgr.reply_to_comment(id1, "Reply to Alice", author_obj("Charlie"), "C")
        mgr.resolve_comment(id2)

        # Save
        output_path = tmp_path / "test_roundtrip.docx"
        doc.save(str(output_path))

        # Reload
        doc2 = Document(str(output_path))
        mgr2 = CommentManager(doc2)

        # Verify comments
        comments = list(mgr2.list_comments())
        assert len(comments) == 3

        # Verify threads
        threads = mgr2.get_comment_threads()
        assert len(threads) == 2

        # Find thread with reply
        thread_with_reply = next(t for t in threads if t.reply_count > 0)
        assert thread_with_reply.root.text == "Comment on first para"
        assert thread_with_reply.replies[0].text == "Reply to Alice"

        # Find resolved thread
        resolved_thread = next(t for t in threads if t.root.text == "Comment on second para")
        assert resolved_thread.is_resolved


class TestPeopleXml:
    """Tests for people.xml integration."""

    def test_people_xml_not_created_by_default(self, tmp_path):
        """people.xml should not be created unless requested."""
        from zipfile import ZipFile

        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "Comment", author_obj("Author"))

        output_path = tmp_path / "no_people.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            names = zf.namelist()
            assert "word/people.xml" not in names

    def test_add_person_without_presence(self, tmp_path):
        """Creating a person entry without presenceInfo should be valid."""
        from lxml import etree
        from zipfile import ZipFile

        doc = Document()
        mgr = CommentManager(doc)
        mgr.ensure_person("Alice")

        output_path = tmp_path / "people_no_presence.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            xml = etree.fromstring(zf.read("word/people.xml"))

        ns_w15 = "http://schemas.microsoft.com/office/word/2012/wordml"
        people = xml.findall(f"{{{ns_w15}}}person")
        assert len(people) == 1
        assert people[0].get(f"{{{ns_w15}}}author") == "Alice"
        assert people[0].find(f"{{{ns_w15}}}presenceInfo") is None

    def test_get_person_helper(self):
        """get_person should raise when missing and return PersonInfo when present."""
        doc = Document()
        mgr = CommentManager(doc)
        with pytest.raises(KeyError):
            mgr.get_person("Alice")

        mgr.ensure_person("Alice")
        person = mgr.get_person("Alice")
        assert isinstance(person, PersonInfo)
        assert person.author == "Alice"

    def test_get_default_author_person_from_docx(self, tmp_path):
        """Resolve default author from a DOCX source."""
        source_doc = Document()
        source_doc.core_properties.author = "Alice"
        source_mgr = CommentManager(source_doc)
        source_mgr.ensure_person(
            "Alice",
            presence={"provider_id": "AD", "user_id": "user"},
        )

        source_path = tmp_path / "author_source.docx"
        source_doc.save(str(source_path))

        mgr = CommentManager(Document())
        person, initials = mgr.get_default_author_person(
            docx_path=str(source_path),
            include_presence=True,
            strict_docx=True,
        )

        assert person.author == "Alice"
        assert person.provider_id == "AD"
        assert person.user_id == "user"
        assert initials is None

    def test_get_default_author_person_strict_docx_missing(self):
        """Strict DOCX mode should raise when the file cannot be read."""
        mgr = CommentManager(Document())
        with pytest.raises(ValueError):
            mgr.get_default_author_person(
                docx_path="does-not-exist.docx",
                strict_docx=True,
            )

    def test_get_default_author_person_warns_on_multiple_people(self, tmp_path, monkeypatch):
        """Multiple people entries should warn and fall back."""
        import warnings
        import docx_comments.system_author as system_author

        source_doc = Document()
        source_mgr = CommentManager(source_doc)
        source_mgr.ensure_person("Alice")
        source_mgr.ensure_person("Bob")

        source_path = tmp_path / "author_source_multi.docx"
        source_doc.save(str(source_path))

        target_doc = Document()
        target_doc.core_properties.author = "Fallback"
        target_mgr = CommentManager(target_doc)

        monkeypatch.setattr(system_author, "_system_office_user_info", lambda: (None, None))

        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            person, _ = target_mgr.get_default_author_person(docx_path=str(source_path))

        assert any("people entries" in str(w.message) for w in caught)
        assert person.author == "Fallback"

    def test_existing_people_xml_preserved_on_comment(self, tmp_path):
        """Adding a comment should not remove existing people.xml data."""
        from lxml import etree
        from zipfile import ZipFile

        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)
        mgr.ensure_person("Alice")

        mgr.add_comment(para, "Comment", author_obj("Alice"))

        output_path = tmp_path / "people_preserved.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            xml = etree.fromstring(zf.read("word/people.xml"))

        ns_w15 = "http://schemas.microsoft.com/office/word/2012/wordml"
        people = xml.findall(f"{{{ns_w15}}}person")
        authors = {p.get(f"{{{ns_w15}}}author") for p in people}
        assert "Alice" in authors

    def test_merge_people_from_document(self):
        """Merging people should add missing authors without overwriting."""
        source_doc = Document()
        source_mgr = CommentManager(source_doc)
        source_mgr.ensure_person(
            "Alice",
            presence={"provider_id": "provider", "user_id": "user"},
        )
        source_mgr.ensure_person("Bob")

        target_doc = Document()
        target_mgr = CommentManager(target_doc)
        target_mgr.ensure_person("Bob")

        added = target_mgr.merge_people_from(source_doc)
        assert [person.author for person in added] == ["Alice"]

        people = target_mgr.get_people()
        authors = {person.author for person in people}
        assert authors == {"Alice", "Bob"}

        alice = next(person for person in people if person.author == "Alice")
        assert alice.provider_id is None
        assert alice.user_id is None

        target_doc_with_presence = Document()
        target_mgr_with_presence = CommentManager(target_doc_with_presence)
        target_mgr_with_presence.merge_people_from(source_doc, include_presence=True)
        people_with_presence = target_mgr_with_presence.get_people()
        alice_with_presence = next(
            person for person in people_with_presence if person.author == "Alice"
        )
        assert alice_with_presence.provider_id == "provider"
        assert alice_with_presence.user_id == "user"

    def test_add_comment_with_author_personinfo(self):
        """Author can be provided as PersonInfo."""
        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)
        mgr.ensure_person("Alice")

        person = mgr.get_people()[0]
        mgr.add_comment(para, "Comment", author=person)

        comments = list(mgr.list_comments())
        assert comments[0].author == "Alice"

    def test_add_comment_with_author_presence_personinfo(self, tmp_path):
        """Author PersonInfo with presence should create people.xml entry."""
        from lxml import etree
        from zipfile import ZipFile

        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)

        person = PersonInfo(author="Alice", provider_id="provider", user_id="user")
        mgr.add_comment(
            para,
            "Comment",
            author=person,
            person=person,
        )

        output_path = tmp_path / "author_dict_presence.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            xml = etree.fromstring(zf.read("word/people.xml"))

        ns_w15 = "http://schemas.microsoft.com/office/word/2012/wordml"
        person = xml.find(f"{{{ns_w15}}}person")
        assert person is not None
        assert person.get(f"{{{ns_w15}}}author") == "Alice"
        presence = person.find(f"{{{ns_w15}}}presenceInfo")
        assert presence is not None
        assert presence.get(f"{{{ns_w15}}}providerId") == "provider"
        assert presence.get(f"{{{ns_w15}}}userId") == "user"
